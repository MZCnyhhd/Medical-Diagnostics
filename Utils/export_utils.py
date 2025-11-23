import io
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
import textwrap

try:
    from docx import Document
except ImportError:
    Document = None

def generate_pdf(content: str, filename: str = "report.pdf") -> io.BytesIO:
    """生成 PDF 文件流"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    # 尝试注册中文字体 (Windows 默认字体)
    try:
        # 常见 Windows 中文字体路径
        font_path = "C:/Windows/Fonts/simhei.ttf"
        pdfmetrics.registerFont(TTFont('SimHei', font_path))
        c.setFont("SimHei", 12)
    except Exception:
        # 如果找不到字体，回退到默认字体（不支持中文）
        c.setFont("Helvetica", 12)

    y = height - 50
    margin = 50
    max_width = width - 2 * margin
    
    # 简单的文本换行处理
    lines = content.split('\n')
    for line in lines:
        # 如果行太长，进行折行
        # 这里只是一个简单的估算，实际应该用 stringWidth 计算
        wrapped_lines = textwrap.wrap(line, width=40) # 40 chars approx
        if not wrapped_lines:
            y -= 15
            continue
            
        for wrapped_line in wrapped_lines:
            if y < 50:
                c.showPage()
                y = height - 50
                try:
                    c.setFont("SimHei", 12)
                except:
                    c.setFont("Helvetica", 12)
            
            c.drawString(margin, y, wrapped_line)
            y -= 15
            
    c.save()
    buffer.seek(0)
    return buffer

def generate_docx(content: str, filename: str = "report.docx") -> io.BytesIO:
    """生成 Word 文件流"""
    if Document is None:
        raise ImportError("python-docx not installed")
        
    doc = Document()
    doc.add_heading('医疗诊断报告', 0)
    
    for line in content.split('\n'):
        doc.add_paragraph(line)
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
